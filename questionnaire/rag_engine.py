"""
RAG Engine - Core AI component for document processing and answer generation
Uses OpenAI embeddings + cosine similarity for retrieval
"""
import json
import math
import re
from typing import List, Tuple, Optional
import PyPDF2
import docx


def _clean_pdf_text(text: str) -> str:
    """
    Fix common PyPDF2 artefacts in extracted PDF text:
    - Soft-hyphen line breaks:  "Multi -\nFactor"  →  "Multi-Factor"
    - Inline spaced hyphens:    "Multi -Factor"      →  "Multi-Factor"
    - Trailing whitespace before newlines
    - Long underscore runs (Source/Answer lines) collapsed to a marker
    """
    # "word -\nnext", "word -\n next", "word-\nnext"  →  "word-next"
    # Also handles trailing " -" at end of line followed by word on next line
    text = re.sub(r' *-\n +', '-', text)   # hyphen at line end joins to next line
    text = re.sub(r'(\w) -\n(\w)', r'\1-\2', text)  # spaced hyphen across lines
    # "word - next" or "word -next" (space around hyphen mid-word)  →  "word-next"
    text = re.sub(r'(\w) - (\w)', r'\1-\2', text)
    text = re.sub(r'(\w) -(\w)', r'\1-\2', text)
    text = re.sub(r'(\w)- (\w)', r'\1-\2', text)
    # Strip trailing spaces from every line
    text = re.sub(r'[ \t]+\n', '\n', text)
    # Collapse long underscore runs
    text = re.sub(r'_{3,}', '___', text)
    return text


def extract_text_from_pdf(file_path: str) -> List[dict]:
    """Extract text with page numbers from PDF, with artefact cleanup."""
    pages = []
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ''
                text = _clean_pdf_text(text)
                if text.strip():
                    pages.append({'text': text, 'page': i + 1})
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return pages


def extract_text_from_docx(file_path: str) -> List[dict]:
    """
    Extract text from DOCX for use as a REFERENCE document (chunking/embedding).
    Walks paragraphs AND tables in body order.

    Important: rows whose first cell looks like a Q-label (Q1, Q2 …) are
    intentionally skipped here — those are questionnaire question rows and are
    handled exclusively by _extract_questions_from_docx_tables.  Emitting them
    here would poison the fallback regex in extract_questions_from_file.
    """
    # Pattern used to detect a Q-label cell so we can skip questionnaire rows
    _q_label_re = re.compile(r'^Q\s*\d+\.?$', re.IGNORECASE)
    # Cells that are just answer/source placeholders — also skip
    _placeholder_re = re.compile(r'^(answer|source|citation|_+|\s*)$', re.IGNORECASE)

    sections = []
    try:
        doc = docx.Document(file_path)

        current_section = ''
        buffer = []

        def flush():
            nonlocal buffer
            if buffer:
                sections.append({'text': '\n'.join(buffer), 'section': current_section})
                buffer = []

        def collect_cell_text(cell) -> str:
            parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            return ' '.join(parts).strip()

        for child in doc.element.body:
            tag = child.tag.split('}')[-1]

            if tag == 'p':
                para = docx.text.paragraph.Paragraph(child, doc)
                style_name = ''
                try:
                    style_name = para.style.name or ''
                except Exception:
                    pass
                text = para.text.strip()
                if style_name.startswith('Heading') and text:
                    flush()
                    current_section = text
                elif text:
                    buffer.append(text)

            elif tag == 'tbl':
                table = docx.table.Table(child, doc)
                for row in table.rows:
                    cells = row.cells
                    if not cells:
                        continue

                    # Get first cell text to decide whether to skip this row
                    first_cell_text = collect_cell_text(cells[0])

                    # Skip questionnaire Q-rows and placeholder rows entirely —
                    # they belong to the question extractor, not the text corpus
                    if _q_label_re.match(first_cell_text):
                        continue
                    if _placeholder_re.match(first_cell_text):
                        continue

                    # For all other table rows, join non-empty cells with a space
                    row_parts = []
                    for cell in cells:
                        ct = collect_cell_text(cell)
                        if ct and not _placeholder_re.match(ct):
                            row_parts.append(ct)
                    if row_parts:
                        buffer.append(' '.join(row_parts))

        flush()

    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return sections


def extract_text_from_txt(file_path: str) -> List[dict]:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return [{'text': content, 'section': 'Document'}]
    except Exception as e:
        print(f"TXT extraction error: {e}")
        return []


def _extract_questions_from_pdf(file_path: str) -> List[str]:
    """
    PDF-specific question extractor.
    Handles: multi-line question wrap, page headers fused to Q-lines,
    and inline category words (separated by double-space in this PDF layout).
    """
    pages = extract_text_from_pdf(file_path)
    if not pages:
        return []

    full_text = "\n".join(p["text"] for p in pages)

    # Remove page headers fused onto Q-label lines by PyPDF2.
    # e.g. "...CONFIDENTIAL | Page 2 Q3 Data Security  What TLS..."
    # Split at the Q-label so it starts on its own line.
    full_text = re.sub(
        r"[^\n]*(?:CONFIDENTIAL|Page\s+\d+)\s+(?=Q\s*\d)",
        "\n", full_text, flags=re.IGNORECASE)
    # Remove any remaining standalone header lines
    full_text = re.sub(
        r"^[^\n]*CONFIDENTIAL[^\n]*$", "",
        full_text, flags=re.IGNORECASE | re.MULTILINE)
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)

    # Learn category vocabulary: text between Q-label and first double-space.
    # e.g. "Q1 Data Security  What data..." → category = "Data Security"
    known_categories: set = set()
    for m in re.finditer(r"^Q\s*\d{1,3}\s+(.+?)  ", full_text, re.MULTILINE):
        cat = m.group(1).strip()
        if 1 <= len(cat.split()) <= 4:
            known_categories.add(cat)

    # Primary: Q-label + category up to double-space + question start
    q_dbl_re   = re.compile(r"^Q\s*\d{1,3}\s+.+?  (.+)$",  re.MULTILINE)
    # Fallback: Q-label + everything after it
    q_plain_re = re.compile(r"^Q\s*\d{1,3}\s+(.+)$",        re.MULTILINE)
    # Numbered list fallback
    num_re     = re.compile(r"^\d{1,3}[.)]\s+(.+)$",          re.MULTILINE)
    stop_re    = re.compile(
        r"^(?:Answer\b|Source\b|___|Section\s|DECLARATION|INSTRUCTIONS)",
        re.IGNORECASE)

    questions: list = []
    current_parts: list = []
    in_question = False

    def flush():
        nonlocal current_parts, in_question
        if current_parts:
            joined = " ".join(" ".join(p.split()) for p in current_parts).strip()
            if len(joined) >= 15:
                questions.append(joined)
        current_parts.clear()
        in_question = False

    for raw_line in full_text.split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            continue
        if stop_re.match(stripped):
            flush()
            continue
        if re.match(r"^Section\s+[A-Z]", stripped):
            flush()
            continue

        # Q-label with double-space separator — most reliable, strips category
        m = q_dbl_re.match(stripped)
        if m:
            flush()
            start_txt = m.group(1).strip()
            if start_txt:
                current_parts = [start_txt]
            in_question = True
            continue

        # Q-label without double-space
        m = q_plain_re.match(stripped)
        if m:
            remainder = m.group(1).strip()
            flush()
            # If remainder is only a known category, question is on next line
            if remainder and remainder not in known_categories:
                current_parts = [remainder]
            in_question = True
            continue

        # Numbered list fallback
        m = num_re.match(stripped)
        if m:
            flush()
            current_parts = [m.group(1).strip()]
            in_question = True
            continue

        # Continuation line
        if in_question:
            current_parts.append(stripped)
            if " ".join(current_parts).rstrip().endswith("?"):
                flush()

    flush()
    return questions


def _extract_questions_from_docx_tables(file_path: str) -> List[str]:
    """
    Primary extraction path for DOCX questionnaires that store questions in tables.

    Recognises four common table layouts:
      Layout A  — Q1 | category | question text       (3+ cols, Q-label first)
      Layout B  — Q1 | question text                  (2 cols, Q-label first)
      Layout C  — No. | question text                 (numeric label first)
      Layout D  — Single-col table with full "Q1. question" text in one cell

    Rules:
    - The question cell is the rightmost non-empty, non-placeholder cell.
    - Multi-paragraph cells are fully joined (no truncation).
    - Cells that are answer boxes, source lines, or decorative dividers are ignored.
    - Questions shorter than 10 chars are discarded as false positives.
    """
    questions: List[str] = []
    try:
        doc = docx.Document(file_path)

        # Matches: Q1  Q1.  Q 1  Q-1  Q01  etc.
        q_label_re = re.compile(r'^Q[\s\-]?\d+\.?$', re.IGNORECASE)
        # Matches: 1.  1)  1:  (01.  etc.)
        num_label_re = re.compile(r'^\d{1,3}[\.\)\:]$')
        # Cells to ignore as question candidates
        placeholder_re = re.compile(
            r'^(answer|source|citation|response|comments?|_+|[-–—]+|\s*)$',
            re.IGNORECASE
        )
        # Single-cell "Q1. full question text" pattern
        inline_q_re = re.compile(
            r'^(?:Q[\s\-]?\d+|Question\s+\d+|No\.?\s*\d+)[\.\:\)]\s+(.{10,})',
            re.IGNORECASE
        )

        def cell_full_text(cell) -> str:
            """Join all paragraphs in a cell, preserving internal spacing."""
            parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            return ' '.join(parts).strip()

        for table in doc.tables:
            if not table.rows:
                continue

            row0_cells = table.rows[0].cells
            if not row0_cells:
                continue

            cell_texts = [cell_full_text(c) for c in row0_cells]
            # Remove duplicate adjacent cells (DOCX sometimes repeats merged cells)
            seen_cells: list = []
            for ct in cell_texts:
                if not seen_cells or ct != seen_cells[-1]:
                    seen_cells.append(ct)
            cell_texts = seen_cells

            first = cell_texts[0] if cell_texts else ''

            # ── Layout D: single cell containing "Q1. Full question text" ──
            if len(cell_texts) == 1:
                m = inline_q_re.match(first)
                if m:
                    questions.append(m.group(1).strip())
                continue

            # ── Layout A/B: first cell is a Q-label ──
            if q_label_re.match(first) or num_label_re.match(first):
                # Walk remaining cells right-to-left, pick first real question text
                question_text = ''
                for ct in reversed(cell_texts[1:]):
                    if ct and not placeholder_re.match(ct) and len(ct) > 10:
                        question_text = ct
                        break
                if question_text:
                    questions.append(question_text)
                continue

            # ── Layout E: first cell is "Q1. Full question text" (multi-col) ──
            m = inline_q_re.match(first)
            if m:
                questions.append(m.group(1).strip())

    except Exception as e:
        print(f"Table question extraction error: {e}")

    return questions


def extract_questions_from_file(file_path: str, file_type: str) -> List[str]:
    """
    Extract questions from a questionnaire file.

    Extraction strategy (tried in order, stops at first success):

    1. PDF   — dedicated PDF extractor that joins all pages into one blob,
               strips running headers, consumes inline category words, and
               accumulates multi-line question text.
    2. DOCX  — table-aware extractor for Q-labelled table rows.
    3. Plain-text line parser (all formats) — numbered/lettered lists with
               multi-line joining and tab-row normalisation.
    """
    questions: List[str] = []

    # ── Path 1: PDF-specific extractor ────────────────────────────────────
    if file_type == 'pdf':
        questions = _extract_questions_from_pdf(file_path)
        if questions:
            print(f"[extract_questions] pdf path: {len(questions)} questions found")
            return _deduplicate(questions)
        print("[extract_questions] pdf path returned 0, falling back to text")

    # ── Path 2: DOCX table scan ────────────────────────────────────────────
    if file_type in ('docx', 'doc'):
        questions = _extract_questions_from_docx_tables(file_path)
        if questions:
            print(f"[extract_questions] table path: {len(questions)} questions found")
            return _deduplicate(questions)
        print("[extract_questions] table path returned 0, falling back to text")

    # ── Path 3: Full-text line parser ──────────────────────────────────────
    if file_type == 'pdf':
        pages = extract_text_from_pdf(file_path)
        full_text = '\n'.join(p['text'] for p in pages)
    elif file_type in ('docx', 'doc'):
        sections = extract_text_from_docx(file_path)
        full_text = '\n'.join(s['text'] for s in sections)
    else:
        sections = extract_text_from_txt(file_path)
        full_text = '\n'.join(s['text'] for s in sections)

    # Pre-process: expand tab-separated rows that look like
    # "Q1\tCategory\tActual question text?" into just the question fragment.
    # This handles the body-walker output for edge-case DOCX files.
    _tab_q_re = re.compile(
        r'^(?:Q[\s\-]?\d+\.?|Question\s+\d+[\.\:\)]|\d{1,3}[\.\)\:])\t',
        re.IGNORECASE
    )
    cleaned_lines = []
    for raw in full_text.split('\n'):
        if _tab_q_re.match(raw):
            # It's a tab-joined Q-table row — extract the last tab column
            parts = [p.strip() for p in raw.split('\t') if p.strip()]
            # Skip Q-label and category-ish short strings; take the longest part
            candidates = [p for p in parts[1:] if len(p) > 15]
            if candidates:
                cleaned_lines.append(max(candidates, key=len))
            # else drop the row — it's just a label/placeholder row
        else:
            cleaned_lines.append(raw)
    full_text = '\n'.join(cleaned_lines)

    # Regex patterns that mark the start of a question line
    start_patterns = [
        re.compile(r'^\d{1,3}[\.\)]\s+\S'),                      # "1. text"  "1) text"
        re.compile(r'^Q[\s\-]?\d+[\.\:\)]\s+\S', re.I),          # "Q1. text"  "Q 1: text"
        re.compile(r'^Question\s+\d+[\.\:\)]\s+\S', re.I),       # "Question 1: text"
        re.compile(r'^No\.?\s*\d+[\.\:\)]\s+\S', re.I),          # "No. 1. text"
    ]
    ends_with_q = re.compile(r'.{15,}\?\s*$')
    structural_re = re.compile(
        r'^(?:[A-Z][A-Z\s]{4,}|Section\s+\w|SECTION\s+\w|Part\s+[A-Z\d])',
    )

    current_parts: List[str] = []

    def flush_question():
        nonlocal current_parts
        if current_parts:
            full = ' '.join(' '.join(p.split()) for p in current_parts)
            if len(full) > 10:
                questions.append(full.strip())
        current_parts = []

    for raw_line in full_text.split('\n'):
        line = raw_line.strip()
        if not line:
            flush_question()
            continue

        matched = False
        for pat in start_patterns:
            if pat.match(line):
                flush_question()
                clean = re.sub(
                    r'^(?:\d{1,3}[\.\)]|Q[\s\-]?\d+[\.\:\)]|Question\s+\d+[\.\:\)]|No\.?\s*\d+[\.\:\)])\s*',
                    '', line, flags=re.IGNORECASE
                ).strip()
                if len(clean) > 5:
                    current_parts = [clean]
                matched = True
                break

        if not matched:
            if current_parts:
                if structural_re.match(line):
                    flush_question()
                else:
                    current_parts.append(line)
                    combined = ' '.join(current_parts)
                    if combined.rstrip().endswith('?'):
                        flush_question()
            elif ends_with_q.match(line):
                flush_question()
                questions.append(line.strip())

    flush_question()

    print(f"[extract_questions] text path: {len(questions)} questions found")
    return _deduplicate(questions)


def _deduplicate(questions: List[str]) -> List[str]:
    """Deduplicate a list preserving order, case-insensitive."""
    seen: set = set()
    unique: List[str] = []
    for q in questions:
        key = q.lower().strip()
        if key not in seen and key:
            seen.add(key)
            unique.append(q)
    return unique


def chunk_text(text: str, chunk_size: int = 600, overlap: int = 100) -> List[str]:
    """Split text into overlapping chunks for better retrieval"""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunks.append(' '.join(chunk_words))
        i += chunk_size - overlap
    return [c for c in chunks if len(c.strip()) > 50]


def get_embeddings(texts: List[str], api_key: str) -> List[List[float]]:
    """Get embeddings using OpenAI API"""
    import urllib.request
    import urllib.error
    
    embeddings = []
    batch_size = 20
    
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        payload = json.dumps({
            'model': 'text-embedding-3-small',
            'input': batch
        }).encode('utf-8')
        
        req = urllib.request.Request(
            'https://api.openai.com/v1/embeddings',
            data=payload,
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            }
        )
        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read())
                for item in data['data']:
                    embeddings.append(item['embedding'])
        except Exception as e:
            print(f"Embedding error: {e}")
            # Return zero vectors as fallback
            for _ in batch:
                embeddings.append([0.0] * 1536)
    
    return embeddings


def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Compute cosine similarity between two vectors"""
    if not vec1 or not vec2:
        return 0.0
    dot = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = math.sqrt(sum(a * a for a in vec1))
    norm2 = math.sqrt(sum(b * b for b in vec2))
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return dot / (norm1 * norm2)


def retrieve_relevant_chunks(
    question: str,
    chunks_with_embeddings: List[Tuple],
    api_key: str,
    top_k: int = 5
) -> List[Tuple]:
    """Retrieve most relevant chunks for a question using semantic search"""
    q_embedding = get_embeddings([question], api_key)[0]
    
    scored = []
    for chunk_id, chunk_text, embedding, doc_name, page_num in chunks_with_embeddings:
        if embedding:
            sim = cosine_similarity(q_embedding, embedding)
            scored.append((sim, chunk_id, chunk_text, doc_name, page_num))
    
    scored.sort(key=lambda x: -x[0])
    return scored[:top_k]


def generate_answer(
    question: str,
    relevant_chunks: List[Tuple],
    api_key: str,
    project_name: str = ""
) -> dict:
    """Generate answer using GPT-4o with retrieved context"""
    import urllib.request
    
    if not relevant_chunks:
        return {
            'answer': 'No relevant information found in the reference documents.',
            'citations': [],
            'confidence': 0.0
        }
    
    # Build context with source attribution
    context_parts = []
    for i, (score, chunk_id, chunk_text, doc_name, page_num) in enumerate(relevant_chunks):
        page_ref = f", Page {page_num}" if page_num else ""
        context_parts.append(f"[Source {i+1}: {doc_name}{page_ref}]\n{chunk_text}")
    
    context = '\n\n---\n\n'.join(context_parts)
    
    system_prompt = """You are a precise security compliance analyst. Your job is to answer vendor security questionnaire questions using ONLY the provided reference documents.

Rules:
1. Answer solely from the provided sources — do not use outside knowledge
2. After each factual claim, add an inline citation: [Source 1], [Source 2], etc.
3. Be thorough: if multiple sources cover different aspects, include all relevant details
4. If a specific detail is not in the sources, state: "Not specified in the provided documents."
5. Write in professional, direct language suitable for a security questionnaire response
6. At the very end of your response, on a new line, write your confidence score in this exact format:
   CONFIDENCE: <number between 0 and 100>
   Where 100 = the sources fully and explicitly answer the question,
         70  = the sources partially answer the question,
         40  = the sources have related but indirect information,
         10  = the sources barely address the question."""

    user_prompt = f"""VENDOR QUESTIONNAIRE QUESTION:
{question}

REFERENCE DOCUMENT EXCERPTS:
{context}

Instructions: Answer the question above using only the source excerpts. Add [Source N] citations inline. End your response with CONFIDENCE: <0-100>."""

    payload = json.dumps({
        'model': 'gpt-4o-mini',
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt}
        ],
        'temperature': 0.1,
        'max_tokens': 800
    }).encode('utf-8')
    
    req = urllib.request.Request(
        'https://api.openai.com/v1/chat/completions',
        data=payload,
        headers={
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
    )
    
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
            full_answer = data['choices'][0]['message']['content']
            
            # Extract confidence score
            confidence = 0.5
            conf_match = re.search(r'CONFIDENCE:\s*\[?(\d+)\]?', full_answer, re.IGNORECASE)
            if conf_match:
                confidence = int(conf_match.group(1)) / 100
                full_answer = full_answer[:conf_match.start()].strip()
            
            # Build citations list
            citations = []
            for i, (score, chunk_id, chunk_text, doc_name, page_num) in enumerate(relevant_chunks):
                if f'[Source {i+1}]' in full_answer:
                    page_ref = f"p.{page_num}" if page_num else ""
                    citations.append({
                        'number': i + 1,
                        'document': doc_name,
                        'page': page_num,
                        'relevance_score': round(score, 3),
                        'excerpt': chunk_text[:200] + '...' if len(chunk_text) > 200 else chunk_text
                    })
            
            return {
                'answer': full_answer,
                'citations': citations,
                'confidence': confidence
            }
    except Exception as e:
        print(f"Generation error: {e}")
        return {
            'answer': f'Error generating answer: {str(e)}',
            'citations': [],
            'confidence': 0.0
        }


def categorize_questions(questions: List[str], api_key: str) -> List[str]:
    """Auto-categorize questions using AI for better organization"""
    import urllib.request
    
    if not questions or not api_key:
        return ['General'] * len(questions)
    
    questions_text = '\n'.join([f"{i+1}. {q}" for i, q in enumerate(questions[:50])])
    
    payload = json.dumps({
        'model': 'gpt-4o-mini',
        'messages': [{
            'role': 'user',
            'content': f"""Categorize each question into a short category (2-4 words max).
Return ONLY a JSON array of category strings, one per question, in order.

Questions:
{questions_text}

Return format: ["Category 1", "Category 2", ...]"""
        }],
        'temperature': 0.2,
        'max_tokens': 500
    }).encode('utf-8')
    
    req = urllib.request.Request(
        'https://api.openai.com/v1/chat/completions',
        data=payload,
        headers={
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
    )
    
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())
            content = data['choices'][0]['message']['content']
            # Extract JSON array
            match = re.search(r'\[.*\]', content, re.DOTALL)
            if match:
                cats = json.loads(match.group())
                # Pad or trim to match questions count
                while len(cats) < len(questions):
                    cats.append('General')
                return cats[:len(questions)]
    except Exception as e:
        print(f"Categorization error: {e}")
    
    return ['General'] * len(questions)